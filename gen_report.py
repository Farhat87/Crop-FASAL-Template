import argparse
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import psycopg2
from collections import defaultdict

DB_CONFIG = {
    "dbname": "TIH_Report",
    "user": "postgres",
    "password": "Farhat@2025",
    "host": "localhost",
    "port": "5432"
}

def fetch_data():
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT c.crop_name, s.state_name, cy.year, cy.method, cy.yield_value
        FROM crop_yields cy
        JOIN crops c ON c.crop_id = cy.crop_id
        JOIN states s ON s.state_id = cy.state_id
        ORDER BY c.crop_name, s.state_name, cy.year, cy.method
    """)
    data = defaultdict(lambda: defaultdict(lambda: defaultdict(dict)))
    for crop, state, year, method, yield_val in cur.fetchall():
        data[crop][state][year][method] = yield_val
    cur.close()
    conn.close()
    return data

def add_footer(doc):
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run("© 2025 IDEAS-TIH. All rights reserved. | ")
    footer_run.add_text(f"Date: {datetime.now().strftime('%d-%m-%Y')} | ")

    page_field = OxmlElement('w:fldSimple')
    page_field.set(qn('w:instr'), 'PAGE')
    page_run = OxmlElement('w:r')
    page_text = OxmlElement('w:t')
    page_text.text = "Page "
    page_run.append(page_text)
    page_field.append(page_run)
    footer_paragraph._element.append(page_field)

def create_report(template_path, output_path, orientation, logo_path):
    data = fetch_data()
    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
        if orientation.upper() == 'LANDSCAPE':
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.1))

    title = doc.add_paragraph("IDEAS - Institute of Data Engineering, Analytics and Science Foundation")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(12)
    title.runs[0].bold = True

    subtitle = doc.add_paragraph("ISI Kolkata | https://www.ideas-tih.org | +91 6289351800")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    hr = doc.add_paragraph()
    p = hr._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("Crop wise Yield forecasts")
    heading_run.bold = True
    heading_run.font.size = Pt(14)

    for crop, states in data.items():
        crop_para = doc.add_paragraph()
        crop_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crop_run = crop_para.add_run(crop.upper())
        crop_run.bold = True
        crop_run.underline = True
        crop_run.font.size = Pt(12)
        crop_run.font.color.rgb = RGBColor(0, 102, 204)

        table = doc.add_table(rows=2, cols=9)
        table.style = "Table Grid"
        headers = [
            "State", "2024-25 ARIMA", "ARIMA RMSE",
            "2024-25 XGBoost", "XGBoost RMSE",
            "2024-25 RF", "RF RMSE",
            "2023-24 MoA&FW", "2022-23 MoA&FW"
        ]
        sub_headers = [
            "", "(Yield)", "(RMSE)",
            "(Yield)", "(RMSE)",
            "(Yield)", "(RMSE)",
            "", ""
        ]
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
        for i, text in enumerate(sub_headers):
            table.rows[1].cells[i].text = text

        for state, years in states.items():
            row = table.add_row().cells
            row[0].text = state
            row[1].text = str(years.get(2024, {}).get("ARIMA", ""))
            row[2].text = str(years.get(2024, {}).get("ARIMA_RMSE", ""))
            row[3].text = str(years.get(2024, {}).get("XGBoost", ""))
            row[4].text = str(years.get(2024, {}).get("XGBoost_RMSE", ""))
            row[5].text = str(years.get(2024, {}).get("Random Forest", ""))
            row[6].text = str(years.get(2024, {}).get("RF_RMSE", ""))
            row[7].text = str(years.get(2023, {}).get("MoA&FW", ""))
            row[8].text = str(years.get(2022, {}).get("MoA&FW", ""))

        doc.add_paragraph()

    add_footer(doc)
    doc.save(output_path)
    print(f"✔️ Report saved to {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate crop yield report.")
    parser.add_argument("-t", "--template", required=True, help="Template input path (not used but required)")
    parser.add_argument("-o", "--output", required=True, help="Output .docx file name")
    parser.add_argument("-f", "--format", choices=["LANDSCAPE", "PORTRAIT"], default="PORTRAIT", help="Page orientation")
    parser.add_argument("-l", "--logo", required=True, help="Path to logo image")

    args = parser.parse_args()
    create_report(args.template, args.output, args.format, args.logo)
